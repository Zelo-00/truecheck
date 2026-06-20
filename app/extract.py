"""Извлечение текста и структуры из документа.

Чистый модуль: без сети и ИИ. Тяжёлые парсеры (docx/pdf) импортируются лениво,
поэтому txt-путь и тесты ядра работают на голой стандартной библиотеке.

Делает три вещи:
  1. read_document(...)     — байты/файл → плоский текст;
  2. split_bibliography(...) — отделить тело от списка литературы и распарсить его;
  3. parse_citations / split_episodes — найти маркеры ссылок и нарезать эпизоды.
"""
from __future__ import annotations

import re

from .models import Citation, Episode, SourceRef

# ----------------------------------------------------------------------------- #
#  1. Чтение документа в плоский текст
# ----------------------------------------------------------------------------- #

# все текстовые форматы, которые принимает загрузчик (для UI/валидации)
SUPPORTED_EXTS = (
    ".txt", ".text", ".md", ".markdown", ".csv", ".tsv", ".log",
    ".html", ".htm", ".xml", ".rtf", ".odt", ".docx", ".doc", ".pdf", ".epub",
)


def read_document(data: bytes, filename: str) -> str:
    """Достаёт плоский текст из любого поддерживаемого текстового формата.

    Диспетчеризация по расширению; тяжёлые парсеры импортируются лениво.
    Неизвестные/повреждённые форматы деградируют до декодирования как текст.
    """
    name = (filename or "").lower()
    try:
        if name.endswith((".txt", ".text", ".md", ".markdown", ".csv", ".tsv", ".log")):
            return _decode(data)
        if name.endswith((".html", ".htm", ".xml")):
            return _read_html(data)
        if name.endswith(".rtf"):
            return _read_rtf(data)
        if name.endswith(".odt"):
            return _read_odt(data)
        if name.endswith(".epub"):
            return _read_epub(data)
        if name.endswith(".docx"):
            return _read_docx(data)
        if name.endswith(".doc"):
            return _read_doc(data)
        if name.endswith(".pdf"):
            return _read_pdf(data)
    except Exception:
        pass  # любой сбой парсера → фолбэк на сырой текст, без краха
    # неизвестное расширение или ошибка парсинга — пробуем как текст
    return _decode(data)


def html_to_text(html: str) -> str:
    """HTML → плоский текст. trafilatura, если есть; иначе аккуратный strip-тегов."""
    try:
        import trafilatura  # ленивый импорт
        extracted = trafilatura.extract(html, include_comments=False,
                                        include_tables=True, favor_recall=True)
        if extracted and extracted.strip():
            return extracted
    except Exception:
        pass
    text = re.sub(r"(?is)<(script|style|head)[^>]*>.*?</\1>", " ", html)
    text = re.sub(r"(?i)</(p|div|li|h[1-6]|br|tr)>", "\n", text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = (text.replace("&nbsp;", " ").replace("&amp;", "&")
            .replace("&lt;", "<").replace("&gt;", ">").replace("&quot;", '"'))
    return text


def _read_html(data: bytes) -> str:
    return html_to_text(_decode(data))


def _read_rtf(data: bytes) -> str:
    text = _decode(data)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)        # hex-escapes
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", " ", text)        # control words
    text = re.sub(r"[{}]", "", text)                       # groups
    return re.sub(r"[ \t]+", " ", text)


def _zip_guard(zf) -> None:
    """Анти zip-бомба: суммарный распакованный размер не должен превышать лимит."""
    from . import config  # ленивый импорт, чтобы ядро оставалось лёгким
    total = sum(i.file_size for i in zf.infolist())
    cap = config.ZIP_MAX_UNCOMPRESSED_MB * 1024 * 1024
    if total > cap:
        raise ValueError(f"архив распаковывается в {total} байт (> {cap}) — отклонён")


def _read_odt(data: bytes) -> str:
    import io
    import zipfile  # ленивый импорт
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        _zip_guard(z)
        xml = z.read("content.xml").decode("utf-8", errors="replace")
    xml = re.sub(r"(?i)</text:p>|</text:h>", "\n", xml)
    return re.sub(r"<[^>]+>", "", xml)


def _read_epub(data: bytes) -> str:
    import io
    import zipfile  # ленивый импорт
    out: list[str] = []
    with zipfile.ZipFile(io.BytesIO(data)) as z:
        _zip_guard(z)
        for n in z.namelist():
            if n.lower().endswith((".xhtml", ".html", ".htm")):
                out.append(html_to_text(z.read(n).decode("utf-8", errors="replace")))
    return "\n\n".join(out)


def _read_doc(data: bytes) -> str:
    """Старый .doc: вытаскиваем ASCII/UTF-16 фрагменты (грубо, без внешних утилит)."""
    txt = data.decode("latin-1", errors="ignore")
    # печатные последовательности длиной от 4 символов
    chunks = re.findall(r"[ -~Ѐ-ӿ\n\r\t]{4,}", txt)
    return "\n".join(chunks)


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_docx(data: bytes) -> str:
    import io
    import zipfile

    from docx import Document  # ленивый импорт
    with zipfile.ZipFile(io.BytesIO(data)) as z:   # docx — это zip; анти zip-бомба
        _zip_guard(z)
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    # таблицы тоже могут содержать список литературы
    for tbl in doc.tables:
        for row in tbl.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def _read_pdf(data: bytes) -> str:
    import io
    from pypdf import PdfReader  # ленивый импорт
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    text = "\n".join(pages)
    # склейка переносов вида «инфор-\nмация»
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    return text


def normalize(text: str) -> str:
    """Унификация пробелов и переносов, сохранение абзацев."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = text.replace(" ", " ").replace(" ", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ----------------------------------------------------------------------------- #
#  2. Список литературы
# ----------------------------------------------------------------------------- #

_BIB_HEADING = re.compile(
    r"^\s*(?:библиографический\s+список|список\s+(?:использованн[ыо][йх]\s+)?"
    r"(?:источник\w*|литератур\w*)|литература|references|bibliography)\s*:?\s*$",
    re.IGNORECASE | re.MULTILINE,
)

_ENTRY_NUM = re.compile(r"^\s*\[?(\d{1,3})[\].)]\s+(.*)$")


def split_bibliography(text: str) -> tuple[str, list[SourceRef]]:
    """Делит текст на тело и список литературы; возвращает (тело, [SourceRef])."""
    text = normalize(text)
    m = None
    for m in _BIB_HEADING.finditer(text):
        pass  # берём последнее вхождение заголовка
    if not m:
        return text, []
    body = text[: m.start()].strip()
    bib_block = text[m.end():].strip()
    return body, parse_bibliography(bib_block)


def parse_bibliography(block: str) -> list[SourceRef]:
    """Парсит пронумерованный список литературы в SourceRef."""
    # объединяем переносы строк внутри одной записи: новая запись начинается с «N. »
    raw_lines = [ln.strip() for ln in block.split("\n")]
    entries: list[str] = []
    for ln in raw_lines:
        if not ln:
            continue
        if _ENTRY_NUM.match(ln):
            entries.append(ln)
        elif entries:
            entries[-1] += " " + ln
        else:
            entries.append(ln)

    sources: list[SourceRef] = []
    for i, e in enumerate(entries, start=1):
        m = _ENTRY_NUM.match(e)
        key = m.group(1) if m else str(i)
        body = m.group(2) if m else e
        sources.append(_parse_entry(key, body))
    return sources


_YEAR = re.compile(r"\b(1[89]\d{2}|20\d{2})\b")
_URL = re.compile(r"https?://[^\s,;]+")
_DOI = re.compile(r"10\.\d{4,9}/[^\s,;]+", re.IGNORECASE)
_AUTHOR = re.compile(r"^([А-ЯЁ][а-яё]+(?:\s+[А-ЯЁ]\.\s*[А-ЯЁ]?\.?)?|[A-Z][a-z]+(?:,?\s+[A-Z]\.)*)")


def _parse_entry(key: str, body: str) -> SourceRef:
    src = SourceRef(key=key, bib_text=body)
    if (mu := _URL.search(body)):
        src.url = mu.group(0).rstrip(".")
    if (md := _DOI.search(body)):
        src.doi = md.group(0).rstrip(".")
    if (my := _YEAR.search(body)):
        src.year = my.group(1)
    if (ma := _AUTHOR.match(body.strip())):
        src.author = ma.group(1).strip()
    # заголовок: до первого « // », « / », « — » или года
    title = re.split(r"\s//\s|\s/\s|\s—\s|\s-\s", body, maxsplit=1)[0]
    if src.author and title.startswith(src.author):
        title = title[len(src.author):].lstrip(" .")
    src.title = title.strip().strip(".")[:300]
    return src


# ----------------------------------------------------------------------------- #
#  3. Маркеры цитирования и эпизоды
# ----------------------------------------------------------------------------- #

# [12]  [12, с. 44]  [1, 2, 3]  [1; 2]  [5, c. 10-12]
_NUMERIC = re.compile(r"\[((?:\d{1,3})(?:\s*[,;]\s*\d{1,3})*)\s*(?:,\s*[сc]\.?\s*([\d\-–]+))?\]")
# (Иванов, 2021)  (Иванов и Петров, 2020)  (Smith, 2019)  (Иванов 2021)
_AUTHOR_YEAR = re.compile(
    r"\(([А-ЯЁA-Z][а-яёa-z]+(?:\s+(?:и|и\s+др\.?|et\s+al\.?|&|and)\s*[А-ЯЁA-Z]?[а-яёa-z]*)?)"
    r"[,\s]+(1[89]\d{2}|20\d{2})[а-яёa-z]?\)"
)


def parse_citations(text: str) -> list[Citation]:
    """Находит все маркеры цитирования в тексте с их позициями."""
    cites: list[Citation] = []
    for m in _NUMERIC.finditer(text):
        nums = re.findall(r"\d{1,3}", m.group(1))
        page = m.group(2)
        for n in nums:
            cites.append(Citation(raw=m.group(0), kind="numeric", ref_key=n,
                                  page=page, start=m.start(), end=m.end()))
    for m in _AUTHOR_YEAR.finditer(text):
        key = f"{m.group(1).strip()} {m.group(2)}"
        cites.append(Citation(raw=m.group(0), kind="author_year", ref_key=key,
                              start=m.start(), end=m.end()))
    for m in _DOI.finditer(text):
        cites.append(Citation(raw=m.group(0), kind="doi", ref_key=m.group(0).rstrip("."),
                              start=m.start(), end=m.end()))
    for m in _URL.finditer(text):
        if _DOI.search(m.group(0)):
            continue
        cites.append(Citation(raw=m.group(0), kind="url", ref_key=m.group(0).rstrip("."),
                              start=m.start(), end=m.end()))
    cites.sort(key=lambda c: c.start)
    return cites


_SENT_SPLIT = re.compile(r"(?<=[.!?…])\s+(?=[«\"“А-ЯЁA-Z0-9])")
_PH = ""  # временная маска точки (длина сохраняется → позиции не сдвигаются)
_PROTECT = re.compile(r"\[[^\]]*\]|\([^)]*\)")
_ABBR = re.compile(
    r"\b(?:с|c|стр|гг?|тт?|рис|табл|см|вып|изд|ред|пер|кн|ч|пп?|др|вв?|нач|кон)\.|"
    r"\bт\.\s*(?:д|е|п|к|ч)\.", re.IGNORECASE)


def _mask_dots(text: str) -> str:
    """Прячет точки, по которым нельзя резать: внутри [..]/(..) и в сокращениях."""
    text = _PROTECT.sub(lambda m: m.group(0).replace(".", _PH), text)
    text = _ABBR.sub(lambda m: m.group(0).replace(".", _PH), text)
    return text


def _sentences(paragraph: str) -> list[tuple[int, int, str]]:
    """Грубая нарезка абзаца на предложения с абсолютными позициями."""
    masked = _mask_dots(paragraph)
    out: list[tuple[int, int, str]] = []
    pos = 0
    for chunk in _SENT_SPLIT.split(masked):
        restored = chunk.replace(_PH, ".")
        idx = paragraph.find(restored, pos)
        if idx < 0:
            idx = pos
        out.append((idx, idx + len(restored), restored))
        pos = idx + len(restored)
    return out


def split_episodes(body: str) -> list[Episode]:
    """Нарезает тело на эпизоды — предложения, содержащие хотя бы одну ссылку."""
    body = normalize(body)
    episodes: list[Episode] = []
    ep_index = 0
    for para_no, para in enumerate(body.split("\n")):
        para = para.strip()
        if not para:
            continue
        base = body.find(para)
        for s_start, s_end, sent in _sentences(para):
            abs_start = (base if base >= 0 else 0) + s_start
            local = parse_citations(sent)
            if not local:
                continue
            for c in local:  # перенос позиций в координаты документа
                c.start += abs_start
                c.end += abs_start
            episodes.append(Episode(index=ep_index, text=sent.strip(),
                                    citations=local, para=para_no))
            ep_index += 1
    return episodes


def analyze_document(text: str) -> tuple[list[Episode], list[SourceRef]]:
    """Полный разбор: текст → (эпизоды, источники)."""
    body, sources = split_bibliography(text)
    episodes = split_episodes(body)
    return episodes, sources
