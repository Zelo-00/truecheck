"""Сериализация отчёта: JSON / Markdown + сохранение на диск."""
from __future__ import annotations

import json
import os

from . import config
from .models import Report, Status


def to_dict(rep: Report) -> dict:
    d = rep.to_dict()
    # Enum → строковое значение для чистого JSON
    for ep in d["episodes"]:
        v = ep.get("verdict")
        if v and isinstance(v.get("status"), Status):
            v["status"] = v["status"].value
    return d


def to_json(rep: Report) -> str:
    return json.dumps(to_dict(rep), ensure_ascii=False, indent=2, default=str)


def to_markdown(rep: Report) -> str:
    lines = [f"# Отчёт сверки: {rep.doc_name}", "",
             f"- Дата: {rep.created_at}",
             f"- Соответствие источникам: **{rep.score}%**",
             f"- Вероятность «ИИ-генерация / литература не соответствует»: **{rep.ai_generated_likelihood}%**",
             f"- ИИ-сверка: {'включена' if rep.ai_used else 'выключена'}",
             "", f"**Вердикт.** {rep.verdict_text}", "", "## Эпизоды", ""]
    for ep in rep.episodes:
        v = ep.verdict
        lines.append(f"### Эпизод {ep.index + 1} — {v.status.ru}")
        lines.append(f"> {ep.text}")
        refs = ", ".join(c.raw for c in ep.citations)
        lines.append(f"- Ссылки: {refs}")
        if v.explanation:
            lines.append(f"- Объяснение: {v.explanation}")
        if v.quote:
            lines.append(f"- Цитата из источника: «{v.quote}»")
        if v.reasons:
            lines.append(f"- Признаки: {', '.join(v.reasons)}")
        lines.append("")
    lines.append("## Источники")
    for s in rep.sources:
        lines.append(f"- [{s.key}] {s.bib_text}  — {s.origin} ({s.note})")
    return "\n".join(lines)


def to_docx_bytes(rep: Report) -> bytes:
    """Серверная генерация .docx-отчёта (python-docx)."""
    import io

    from docx import Document  # ленивый импорт

    doc = Document()
    doc.add_heading(f"Отчёт сверки: {rep.doc_name}", level=0)
    doc.add_paragraph(f"Дата: {rep.created_at}")
    doc.add_paragraph(f"Соответствие источникам: {rep.score}%")
    doc.add_paragraph(f"Вероятность «ИИ-генерация / литература не соответствует»: "
                      f"{rep.ai_generated_likelihood}%")
    doc.add_paragraph(f"ИИ-сверка: {'включена' if rep.ai_used else 'выключена'}")
    doc.add_paragraph(rep.verdict_text)

    doc.add_heading(f"Эпизоды ({len(rep.episodes)})", level=1)
    for ep in rep.episodes:
        v = ep.verdict
        doc.add_heading(f"Эпизод {ep.index + 1} — {v.status.ru}", level=2)
        doc.add_paragraph(ep.text)
        doc.add_paragraph("Ссылки: " + (", ".join(c.raw for c in ep.citations) or "—"))
        if v.explanation:
            doc.add_paragraph("Объяснение: " + v.explanation)
        if v.quote:
            p = doc.add_paragraph()
            p.add_run(f"Цитата из источника: «{v.quote}»").italic = True

    doc.add_heading("Источники", level=1)
    for s in rep.sources:
        doc.add_paragraph(f"[{s.key}] {s.bib_text}  — {s.origin} ({s.note})")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def save(rep: Report) -> dict:
    """Сохраняет JSON / Markdown / DOCX в REPORTS_DIR. Возвращает пути."""
    config.ensure_dirs()
    base = os.path.join(config.REPORTS_DIR, rep.job_id)
    with open(base + ".json", "w", encoding="utf-8") as f:
        f.write(to_json(rep))
    with open(base + ".md", "w", encoding="utf-8") as f:
        f.write(to_markdown(rep))
    paths = {"json": base + ".json", "md": base + ".md"}
    try:
        with open(base + ".docx", "wb") as f:
            f.write(to_docx_bytes(rep))
        paths["docx"] = base + ".docx"
    except Exception:  # noqa: BLE001 — docx не критичен, отчёт уже сохранён
        pass
    return paths


def _index_path() -> str:
    return os.path.join(config.REPORTS_DIR, "index.jsonl")


def append_index(rep: Report, sid: str = "") -> None:
    """Заносит отчёт в журнал истории с привязкой к сессии (sid) пользователя."""
    entry = {"job_id": rep.job_id, "doc_name": rep.doc_name,
             "created_at": rep.created_at, "score": rep.score,
             "risk": rep.ai_generated_likelihood, "episodes": len(rep.episodes),
             "ai_used": rep.ai_used, "sid": sid}
    config.ensure_dirs()
    with open(_index_path(), "a", encoding="utf-8") as f:
        f.write(json.dumps(entry, ensure_ascii=False) + "\n")


def read_index(sid: str, limit: int = 100) -> list[dict]:
    """История ТОЛЬКО текущей сессии (по sid), новые — первыми. Без sid → пусто."""
    p = _index_path()
    if not sid or not os.path.exists(p):
        return []
    out: list[dict] = []
    with open(p, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                e = json.loads(line)
            except json.JSONDecodeError:
                continue
            if e.get("sid") == sid:
                out.append(e)
    out.reverse()
    return out[:limit]
